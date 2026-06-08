"""
Convert phan_BC.md → phan_BC.docx
- Tất cả text màu đen tuyệt đối (kể cả heading, link)
- Hỗ trợ: heading 1-6, bold, italic, code inline, code block,
          table, blockquote, bullet/numbered list, horizontal rule
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_black(run):
    run.font.color.rgb = BLACK

def set_para_black(para):
    for run in para.runs:
        set_black(run)

def set_cell_black(cell):
    for para in cell.paragraphs:
        set_para_black(para)

def add_heading(doc, text, level):
    """Add heading với màu đen, không có màu theme."""
    # Xóa anchor {#...} nếu có
    text = re.sub(r'\s*\{#[^}]*\}', '', text).strip()
    # Xử lý bold markers còn sót
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)

    # Level mapping: md h1→Word h1, md h2→h2, md h3→h3, etc.
    # Nhưng vì Part B/C dùng h1 cho phần lớn, dùng Custom styles
    style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3',
                 4: 'Heading 4', 5: 'Heading 5', 6: 'Heading 6'}
    style = style_map.get(level, 'Heading 3')

    para = doc.add_heading(text, level=level)
    para.style = doc.styles[style]

    # Ép màu đen và bỏ theme color
    for run in para.runs:
        run.font.color.rgb = BLACK
        run.font.color.theme_color = None

    # Xóa color trong paragraph XML
    pPr = para._p.get_or_add_pPr()
    # Đảm bảo style không có color
    rPr_list = para._p.findall('.//' + qn('w:rPr'))
    for rPr in rPr_list:
        for color_el in rPr.findall(qn('w:color')):
            rPr.remove(color_el)

    return para


def add_inline_runs(para, text):
    """Parse inline markdown: **bold**, *italic*, `code` → runs."""
    # Pattern: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        # Text trước match
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            set_black(run)
        full = m.group(0)
        if full.startswith('**'):
            run = para.add_run(m.group(2))
            run.bold = True
            set_black(run)
        elif full.startswith('*'):
            run = para.add_run(m.group(3))
            run.italic = True
            set_black(run)
        elif full.startswith('`'):
            run = para.add_run(m.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            set_black(run)
        pos = m.end()
    # Phần còn lại
    if pos < len(text):
        run = para.add_run(text[pos:])
        set_black(run)


def add_normal_para(doc, text, style='Normal'):
    """Thêm paragraph thường với inline markdown support."""
    # Bỏ link syntax [text](url) → chỉ giữ text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    para = doc.add_paragraph(style=style)
    add_inline_runs(para, text)
    return para


def add_table(doc, header_row, data_rows, align_row):
    """Tạo bảng Word từ markdown table data."""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = 'Table Grid'

    # Header
    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(header_row):
        cell_text = cell_text.strip()
        cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
        hdr_cells[i].text = cell_text
        # Bold header
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                set_black(run)
        set_cell_black(hdr_cells[i])

    # Data rows
    for r_idx, row in enumerate(data_rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row):
            cell_text = cell_text.strip()
            cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            cell_text = re.sub(r'\*(.+?)\*', r'\1', cell_text)
            cell_text = re.sub(r'`(.+?)`', r'\1', cell_text)
            cells[c_idx].text = cell_text
            set_cell_black(cells[c_idx])

    # Đảm bảo tất cả text trong bảng màu đen
    for row in table.rows:
        for cell in row.cells:
            set_cell_black(cell)

    doc.add_paragraph()  # Khoảng cách sau bảng
    return table


# ─── Markdown Parser ──────────────────────────────────────────────────────────

def parse_table_row(line):
    """'| a | b | c |' → ['a', 'b', 'c']"""
    parts = [p.strip() for p in line.strip().strip('|').split('|')]
    return parts


def is_separator_row(line):
    """Kiểm tra dòng phân cách bảng: | :--- | :---: |"""
    stripped = line.strip().strip('|')
    return bool(re.match(r'^[\s:\-|]+$', stripped))


def convert(md_path: Path, docx_path: Path):
    doc = Document()

    # ── Thiết lập font mặc định ────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.color.rgb = BLACK

    # Thiết lập tất cả heading styles
    for i in range(1, 7):
        hname = f'Heading {i}'
        if hname in doc.styles:
            h = doc.styles[hname]
            h.font.color.rgb = BLACK
            try:
                h.font.color.theme_color = None
            except Exception:
                pass

    # ── Thiết lập margins ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)

    text = md_path.read_text(encoding='utf-8')
    lines = text.splitlines()

    i = 0
    total = len(lines)

    # State cho bảng
    in_table = False
    table_header = []
    table_aligns = []
    table_data = []

    # State cho code block
    in_code = False
    code_lines = []

    # State cho list
    in_list = False
    list_items = []
    list_ordered = False

    def flush_list():
        nonlocal in_list, list_items, list_ordered
        if not in_list:
            return
        for item in list_items:
            style_name = 'List Number' if list_ordered else 'List Bullet'
            try:
                p = doc.add_paragraph(style=style_name)
            except Exception:
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.left_indent = Inches(0.25)
            # Xóa link syntax
            item = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', item)
            add_inline_runs(p, item)
        in_list = False
        list_items = []
        list_ordered = False

    def flush_table():
        nonlocal in_table, table_header, table_aligns, table_data
        if in_table and table_header:
            add_table(doc, table_header, table_data, table_aligns)
        in_table = False
        table_header = []
        table_aligns = []
        table_data = []

    def flush_code():
        nonlocal in_code, code_lines
        if code_lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            code_text = '\n'.join(code_lines)
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            set_black(run)
        in_code = False
        code_lines = []

    while i < total:
        line = lines[i]

        # ── Code block ````
        if line.strip().startswith('```'):
            if in_table: flush_table()
            if in_list: flush_list()
            if in_code:
                flush_code()
            else:
                in_code = True
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Horizontal rule ---
        if re.match(r'^---+\s*$', line.strip()):
            flush_table()
            flush_list()
            doc.add_paragraph('─' * 60, style='Normal')
            i += 1
            continue

        # ── Heading # ## ### etc.
        hm = re.match(r'^(#{1,6})\s+(.*)', line)
        if hm:
            flush_table()
            flush_list()
            level = len(hm.group(1))
            heading_text = hm.group(2).strip()
            add_heading(doc, heading_text, level)
            i += 1
            continue

        # ── Table row
        if line.strip().startswith('|') and '|' in line[1:]:
            flush_list()
            row_parts = parse_table_row(line)
            if not in_table:
                # Đây là header row
                in_table = True
                table_header = row_parts
                table_aligns = []
                table_data = []
            elif is_separator_row(line):
                # Separator row — lưu align info
                table_aligns = row_parts
            else:
                table_data.append(row_parts)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── Blockquote >
        if line.strip().startswith('>'):
            flush_list()
            bq_text = line.strip().lstrip('>').strip()
            bq_text = re.sub(r'\*\*(.+?)\*\*', r'\1', bq_text)
            try:
                p = doc.add_paragraph(style='Quote')
            except Exception:
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(bq_text)
            run.italic = True
            set_black(run)
            i += 1
            continue

        # ── Bullet list * item  hoặc - item
        bm = re.match(r'^(\s*)[\*\-]\s+(.*)', line)
        if bm:
            if not in_list:
                in_list = True
                list_ordered = False
                list_items = []
            list_items.append(bm.group(2).strip())
            i += 1
            continue

        # ── Numbered list 1. item
        nm = re.match(r'^\d+\.\s+(.*)', line)
        if nm:
            if not in_list:
                in_list = True
                list_ordered = True
                list_items = []
            list_items.append(nm.group(1).strip())
            i += 1
            continue
        else:
            if in_list:
                flush_list()

        # ── Empty line
        if line.strip() == '':
            flush_table()
            flush_list()
            i += 1
            continue

        # ── Normal paragraph
        flush_table()
        flush_list()
        # Bỏ anchor {#...} ở cuối dòng
        clean = re.sub(r'\s*\{#[^}]*\}', '', line).strip()
        if clean:
            add_normal_para(doc, clean)

        i += 1

    # Flush tồn đọng
    flush_table()
    flush_list()
    flush_code()

    # ── Đảm bảo TẤT CẢ text trong document màu đen ───────────────────────────
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.color.rgb = BLACK

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = BLACK

    doc.save(str(docx_path))
    print(f"Saved: {docx_path}")
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables:     {len(doc.tables)}")


if __name__ == '__main__':
    base = Path(__file__).parent.parent / 'docs' / 'report'
    md   = base / 'phan_BC.md'
    docx = base / 'phan_BC.docx'
    convert(md, docx)
