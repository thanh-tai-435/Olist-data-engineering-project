"""Convert a single markdown file to DOCX using the project's build_docx styles."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = Pt(20)

    for level, size, italic in [(1, 14, False), (2, 13, False), (3, 13, True)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name  = "Times New Roman"
        h.font.size  = Pt(size)
        h.font.bold  = True
        h.font.italic = italic
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after  = Pt(6)

    return doc


def add_code_block(doc, code: str):
    for line in code.strip().split("\n"):
        p = doc.add_paragraph()
        p.style = "Normal"
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent  = Cm(1)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)


def add_table_from_md(doc, lines: list):
    rows = [l for l in lines if l.strip().startswith("|") and "---" not in l]
    if not rows:
        return
    parsed = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    num_cols = len(parsed[0])
    table = doc.add_table(rows=len(parsed), cols=num_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(parsed):
        for c_idx, cell_text in enumerate(row[:num_cols]):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    if r_idx == 0:
                        run.font.bold = True
    doc.add_paragraph()


def process_inline(text: str):
    parts = []
    pattern = re.compile(r"(\*\*(.+?)\*\*|`(.+?)`|\*(.+?)\*|([^*`]+))")
    for m in pattern.finditer(text):
        if m.group(2):   parts.append((True,  False, False, m.group(2)))
        elif m.group(3): parts.append((False, False, True,  m.group(3)))
        elif m.group(4): parts.append((False, True,  False, m.group(4)))
        elif m.group(5): parts.append((False, False, False, m.group(5)))
    return parts


def add_inline_para(doc, text: str, style: str = "Normal"):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for bold, italic, code, content in process_inline(text):
        run = p.add_run(content)
        run.font.name = "Courier New" if code else "Times New Roman"
        run.font.size = Pt(10) if code else Pt(13)
        run.bold   = bold
        run.italic = italic
    return p


def parse_md(doc, md_text: str):
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1; continue

        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i]); i += 1
            add_code_block(doc, "\n".join(code_lines))
            i += 1; continue

        if line.strip().startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i]); i += 1
            add_table_from_md(doc, tbl); continue

        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:].strip(), level=1); i += 1; continue
        if line.startswith("## ") and not line.startswith("### "):
            doc.add_heading(line[3:].strip(), level=2); i += 1; continue
        if line.startswith("### ") and not line.startswith("#### "):
            doc.add_heading(line[4:].strip(), level=3); i += 1; continue
        if line.startswith("#### "):
            p = doc.add_paragraph()
            run = p.add_run(line[5:].strip())
            run.bold = True; run.font.name = "Times New Roman"; run.font.size = Pt(13)
            i += 1; continue

        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 60); i += 1; continue

        if line.strip().startswith("- ") or line.strip().startswith("* "):
            add_inline_para(doc, line.strip()[2:].strip(), "List Bullet"); i += 1; continue

        if re.match(r"^\d+\. ", line.strip()):
            add_inline_para(doc, re.sub(r"^\d+\. ", "", line.strip()), "List Number"); i += 1; continue

        add_inline_para(doc, line.strip())
        i += 1


def convert(md_path: str):
    src  = Path(md_path)
    dest = src.with_suffix(".docx")
    doc  = setup_document()
    parse_md(doc, src.read_text(encoding="utf-8"))
    doc.save(dest)
    print(f"Saved: {dest}  ({dest.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    target = sys.argv[1] if len(sys.argv) > 1 else "pham_vi_ca_nhan.md"
    convert(target)
