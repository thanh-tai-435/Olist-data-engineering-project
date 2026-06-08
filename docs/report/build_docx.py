"""
Build DOCX report from markdown chapter files.
Usage: python build_docx.py
Output: docs/report/Olist_Lakehouse_Report.docx

Requires: pip install python-docx
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CHAPTER_FILES = [
    "ch1_introduction.md",
    "ch2_theory.md",
    "ch3_design.md",
    "ch4_implementation.md",
    "ch5_evaluation.md",
    "ch6_conclusion.md",
]

REPORT_DIR = Path(__file__).parent
OUTPUT_FILE = REPORT_DIR / "Olist_Lakehouse_Report.docx"


def setup_document() -> Document:
    doc = Document()

    # Page margins (2cm each side)
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # Normal style: Times New Roman 13pt, justified
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(13)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)
    pf.line_spacing = Pt(20)

    # Heading 1 (chapters)
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.italic = True
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)

    return doc


def add_code_block(doc: Document, code: str):
    """Add a styled code block paragraph."""
    lines = code.strip().split("\n")
    for line in lines:
        p = doc.add_paragraph()
        p.style = "Normal"
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        pf = p.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.left_indent = Cm(1)
        # Light gray background via shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)


def add_table_from_md(doc: Document, lines: list[str]):
    """Parse a markdown table and add as Word table."""
    rows = [l for l in lines if l.strip().startswith("|") and "---" not in l]
    if not rows:
        return

    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return

    num_cols = len(parsed[0])
    table = doc.add_table(rows=len(parsed), cols=num_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(parsed):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= num_cols:
                break
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    if r_idx == 0:
                        run.font.bold = True

    doc.add_paragraph()  # Spacing after table


def process_inline(text: str):
    """Return list of (is_bold, is_italic, is_code, content) tuples."""
    # Very simple inline parser: bold (**), italic (*), code (`)
    parts = []
    pattern = re.compile(r"(\*\*(.+?)\*\*|`(.+?)`|\*(.+?)\*|([^*`]+))")
    for m in pattern.finditer(text):
        if m.group(2):
            parts.append((True, False, False, m.group(2)))
        elif m.group(3):
            parts.append((False, False, True, m.group(3)))
        elif m.group(4):
            parts.append((False, True, False, m.group(4)))
        elif m.group(5):
            parts.append((False, False, False, m.group(5)))
    return parts


def add_paragraph_with_inline(doc: Document, text: str, style: str = "Normal"):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for bold, italic, code, content in process_inline(text):
        run = p.add_run(content)
        run.font.name = "Courier New" if code else "Times New Roman"
        run.font.size = Pt(10) if code else Pt(13)
        run.bold = bold
        run.italic = italic
    return p


def parse_and_add_markdown(doc: Document, md_text: str):
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Code block
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        # Table (detect by first pipe)
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table_from_md(doc, table_lines)
            continue

        # Heading 1
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip().strip("*")
            doc.add_heading(text, level=1)
            i += 1
            continue

        # Heading 2
        if line.startswith("## ") and not line.startswith("### "):
            text = line[3:].strip().strip("*")
            doc.add_heading(text, level=2)
            i += 1
            continue

        # Heading 3
        if line.startswith("### "):
            text = line[4:].strip().strip("*")
            doc.add_heading(text, level=3)
            i += 1
            continue

        # Heading 4 (####) → bold paragraph
        if line.startswith("#### "):
            text = line[5:].strip().strip("*")
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(13)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # Blockquote
        if line.startswith("> "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.right_indent = Cm(1)
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.italic = True
            i += 1
            continue

        # Unordered list
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            for bold, italic, code, content in process_inline(text):
                run = p.add_run(content)
                run.font.name = "Courier New" if code else "Times New Roman"
                run.font.size = Pt(10) if code else Pt(13)
                run.bold = bold
                run.italic = italic
            i += 1
            continue

        # Ordered list (simple: starts with digit + dot)
        if re.match(r"^\d+\. ", line.strip()):
            text = re.sub(r"^\d+\. ", "", line.strip())
            p = doc.add_paragraph(style="List Number")
            for bold, italic, code, content in process_inline(text):
                run = p.add_run(content)
                run.font.name = "Courier New" if code else "Times New Roman"
                run.font.size = Pt(10) if code else Pt(13)
                run.bold = bold
                run.italic = italic
            i += 1
            continue

        # Italic-only line (like *Note text*)
        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            text = line.strip().strip("*")
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            i += 1
            continue

        # Normal paragraph
        add_paragraph_with_inline(doc, line.strip())
        i += 1


def add_centered(doc: Document, text: str, size: int = 13,
                 bold: bool = False, italic: bool = False, space_after: int = 6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    return p


def add_info_row(doc: Document, label: str, value: str):
    """Two-column style: label right-aligned, value left — simulated with tabs."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{label}: {value}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    return p


def add_cover_page(doc: Document):
    # ── Logo placeholder (empty bordered box) ──────────────────────────
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_before = Pt(12)
    logo_p.paragraph_format.space_after = Pt(4)
    run = logo_p.add_run("[  LOGO TRƯỜNG  ]")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── Tên trường ──────────────────────────────────────────────────────
    add_centered(doc,
                 "ĐẠI HỌC QUỐC GIA THÀNH PHỐ HỒ CHÍ MINH",
                 size=13, bold=True, space_after=2)
    add_centered(doc,
                 "TRƯỜNG ĐẠI HỌC KINH TẾ - LUẬT",
                 size=14, bold=True, space_after=2)
    add_centered(doc,
                 "KHOA HỆ THỐNG THÔNG TIN",
                 size=13, bold=False, space_after=20)

    # ── Đường kẻ ────────────────────────────────────────────────────────
    add_centered(doc, "─" * 55, size=11, space_after=20)

    # ── Tiêu đề báo cáo ─────────────────────────────────────────────────
    add_centered(doc, "BÁO CÁO ĐỒ ÁN MÔN HỌC",
                 size=14, bold=True, space_after=18)

    add_centered(doc, "XÂY DỰNG NỀN TẢNG UNIFIED LAKEHOUSE",
                 size=16, bold=True, space_after=4)
    add_centered(doc, "CHO PHÂN TÍCH DỮ LIỆU THƯƠNG MẠI ĐIỆN TỬ",
                 size=16, bold=True, space_after=4)
    add_centered(doc, "(Olist Brazilian E-Commerce Dataset)",
                 size=12, italic=True, space_after=20)

    # ── Hướng ───────────────────────────────────────────────────────────
    add_centered(doc, "Hướng: Data Engineering / Lakehouse",
                 size=13, italic=True, space_after=24)

    # ── Thông tin nhân sự ────────────────────────────────────────────────
    add_centered(doc, "─" * 55, size=11, space_after=14)

    info_block = [
        ("Sinh viên thực hiện",  "Phạm Trần Thanh Tài"),
        ("GVHD",                 "Th.S Nguyễn Quang Phúc"),
        ("HD Doanh nghiệp",      "Anh Nguyễn Duy Duẩn"),
    ]
    for label, value in info_block:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        r_label = p.add_run(f"{label}:  ")
        r_label.font.name = "Times New Roman"
        r_label.font.size = Pt(13)
        r_value = p.add_run(value)
        r_value.font.name = "Times New Roman"
        r_value.font.size = Pt(13)
        r_value.bold = True

    # ── Năm ─────────────────────────────────────────────────────────────
    add_centered(doc, "─" * 55, size=11, space_after=18)
    add_centered(doc, "TP. Hồ Chí Minh, 2026",
                 size=13, space_after=0)

    doc.add_page_break()


def main():
    print("Building DOCX report...")
    doc = setup_document()

    add_cover_page(doc)

    for filename in CHAPTER_FILES:
        filepath = REPORT_DIR / filename
        if not filepath.exists():
            print(f"WARNING: {filename} not found, skipping.")
            continue

        print(f"  Processing {filename}...")
        md_text = filepath.read_text(encoding="utf-8")
        parse_and_add_markdown(doc, md_text)
        doc.add_page_break()

    doc.save(OUTPUT_FILE)
    print(f"\nDone! Report saved to: {OUTPUT_FILE}")
    print(f"File size: {OUTPUT_FILE.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
